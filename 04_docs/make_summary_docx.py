# -*- coding: utf-8 -*-
"""하늘하늘중국어 라이브 그룹 클래스 — 한 장짜리 요약본(.docx) · 미니멀 버전"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x7f, 0x00, 0x05)
INK = RGBColor(0x26, 0x26, 0x26)
GRAY = RGBColor(0x80, 0x80, 0x80)
LINE = "d9d9d9"
FONT = "맑은 고딕"


def kfont(run, size=None, bold=None, color=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rfonts.set(qn(a), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = color if color is not None else INK


def para(doc, text="", size=10, bold=False, color=None, before=0, after=3,
         align=None, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.06
    if align: p.alignment = align
    if indent is not None: pf.left_indent = Cm(indent)
    if text:
        kfont(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text):
    p = para(doc, text, size=11, bold=True, color=PRIMARY, before=7, after=3)
    return p


def set_bottom_border(el, color, sz="4"):
    """el = cell._tc ; adds bottom border via tcBorders"""
    tcPr = el.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
    borders.append(b)


def shade(cell, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def cell_text(cell, text, size=9, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    kfont(p.add_run(text), size=size, bold=bold, color=color)


def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left),
                     ('end', right), ('left', left), ('right', right)):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tcPr.append(mar)


def theme_box(doc, items, gap_w=0.5):
    """7월·8월 테마를 색 박스로 나란히 (강조)"""
    cols = len(items) * 2 - 1
    t = doc.add_table(rows=1, cols=cols)
    t.allow_autofit = False
    box_w = (17.2 - gap_w * (len(items) - 1)) / len(items)
    ci = 0
    for idx, (month, theme, desc) in enumerate(items):
        c = t.rows[0].cells[ci]
        c.width = Cm(box_w)
        shade(c, "f4eaea")
        set_cell_margins(c)
        c.text = ""
        p1 = c.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2); p1.paragraph_format.space_before = Pt(2)
        kfont(p1.add_run(month + "  "), size=14, bold=True, color=PRIMARY)
        kfont(p1.add_run(theme), size=11, bold=True, color=INK)
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(2); p2.paragraph_format.space_before = Pt(0)
        kfont(p2.add_run(desc), size=8.5, color=GRAY)
        ci += 1
        if idx < len(items) - 1:  # 가운데 여백 칸
            t.rows[0].cells[ci].width = Cm(gap_w)
            ci += 1
    return t


def clean_table(doc, rows, widths, header=True, header_fill="f4eeee"):
    """가로선만 있는 깔끔한 표 (세로선·격자 없음)"""
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.allow_autofit = False
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            is_head = header and ri == 0
            cell = t.rows[ri].cells[ci]
            cell_text(cell, val, size=9,
                      bold=is_head, color=PRIMARY if is_head else INK)
            cell.width = Cm(widths[ci])
            # 가로선: 헤더는 진한 레드, 본문은 옅은 회색
            set_bottom_border(cell._tc, "7f0005" if is_head else LINE,
                              sz="6" if is_head else "4")
            if is_head:
                shade(cell, header_fill)
    for r in t.rows:
        for ci, w in enumerate(widths):
            r.cells[ci].width = Cm(w)
    return t


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.1); s.bottom_margin = Cm(0.9)
    s.left_margin = Cm(1.7); s.right_margin = Cm(1.7)

# ── 제목 ──
para(doc, "라이브 그룹 클래스 제안 요약", size=18, bold=True, color=PRIMARY,
     after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "하늘하늘중국어", size=10, bold=True, color=GRAY, after=1,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "1:1 케어를 지키며 수입 천장을 올리는 첫 확장",
     size=9.5, color=GRAY, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── 한눈에 ──
heading(doc, "한눈에")
clean_table(doc, [
    ["일정", "매주 수요일 21:00 · 월 4회 · 회당 50~60분"],
    ["가격", "월 5만원"],
    ["형태", "라이브(줌) + 다시보기(유튜브, 10월까지) + 녹음 숙제 1회/월"],
    ["피드백", "녹음 + 코멘트, 제출 후 3~5일 내 (채널톡 1:1)"],
], widths=[2.6, 14.0], header=False)

# ── 등록 ──
heading(doc, "등록 — 7·8월 한 번에 모집, 3가지 중 선택")
clean_table(doc, [
    ["선택", "내용"],
    ["7+8월 패키지", "2개월 묶음 가격 + 가치 추가(1:1 우선권·다시보기 연장·복습자료)"],
    ["7월만 / 8월만", "각 5만원 · 각 월 독립 테마라 단독 수강 OK"],
], widths=[3.4, 13.2])

# ── 테마 (제안, 강조) ──
heading(doc, "테마 제안 — 두 달 커리큘럼")
theme_box(doc, [
    ("7월", "발음·성조 집중", "한국인이 막히는 발음을 한국인 시선으로"),
    ("8월", "드라마 회화 표현", "좋아하는 드라마로 실전 표현 익히기"),
])
para(doc, "각 월 단독 수강 OK · 두 달 함께 들으면 발음 → 표현으로 자연스럽게 이어짐",
     size=8.5, color=GRAY, before=3, after=3, indent=0.1)

# ── 정원 ──
heading(doc, "정원 — 추천 15명")
clean_table(doc, [
    ["정원", "월 매출", "피드백 시간", "월 총투입"],
    ["15명  (추천)", "75만원", "5시간", "약 9시간"],
    ["20명", "100만원", "6.7시간", "약 11시간"],
    ["30명", "150만원", "10시간", "약 14시간"],
    ["40명", "200만원", "13.3시간", "약 17시간"],
], widths=[4.2, 4.2, 4.2, 4.0])
para(doc, "라이브는 인원과 무관하게 4시간 고정 → 정원의 진짜 한계는 ‘개별 피드백 시간’.",
     size=9, color=GRAY, after=3, indent=0.1)

# ── 환불 (표 대신 한 줄) ──
heading(doc, "환불 — 들은 회차만 차감")
para(doc, "첫 수업 전 100%   ·   1회차 후 37,500원   ·   2회차 후 25,000원   ·   3회차 이후 없음",
     size=9.5, after=2, indent=0.1)
para(doc, "패키지의 8월분(미개강)은 전액 환불 · 환불 시 다시보기 종료",
     size=9, color=GRAY, after=3, indent=0.1)

# ── 핵심 원칙 ──
heading(doc, "핵심 원칙")
for line in [
    "그룹 클래스는 1:1 만석 상태에서 신규가 들어올 ‘유일한 문’이자 완결된 목적지.",
    "발화 연습은 녹음 숙제로, 라이브는 강의 + ‘수요일 9시’ 공부 습관.",
    "정원·피드백 기간을 정하는 것이 번아웃 방어선 (“시간 내면 된다” 금물).",
    "시험 없음 · 할인 없음 · 소수정예와 케어 유지.",
]:
    para(doc, "·  " + line, size=9.5, after=2, indent=0.1)

para(doc, "2026-06-01 · 상세 기획서 별도",
     size=8, color=GRAY, before=5, align=WD_ALIGN_PARAGRAPH.RIGHT)

out = r"c:/development/TutorManager_For_Notion/docs/라이브_그룹클래스_요약.docx"
doc.save(out)
print("saved:", out)

# 워드가 설치돼 있으면 PDF까지 자동 생성 (docx2pdf)
try:
    from docx2pdf import convert
    convert(out, r"c:/development/TutorManager_For_Notion/docs/라이브_그룹클래스_요약.pdf")
    print("pdf:", "라이브_그룹클래스_요약.pdf")
except Exception as e:
    print("PDF 변환 생략:", e)
